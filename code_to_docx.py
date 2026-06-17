#!/usr/bin/env python3
"""
Generate a Word document from a directory of code files.

Usage:
    python scripts/code_to_docx.py <directory> [--team TEAM_NAME] [--output FILE.docx]

Example:
    python scripts/code_to_docx.py /tmp/team-42-code --team team-42
    python scripts/code_to_docx.py ./my-project --team team-42 --output submission.docx
"""

import argparse
import json
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


# File extensions to include
CODE_EXTENSIONS = {
    ".py", ".ipynb", ".yaml", ".yml", ".json", ".sh", ".bash",
    ".txt", ".md", ".csv", ".toml", ".cfg", ".ini", ".conf",
    ".html", ".css", ".js", ".ts", ".jsx", ".tsx",
    ".c", ".cpp", ".h", ".hpp", ".java", ".go", ".rs",
    ".r", ".R", ".sql", ".dockerfile", ".env",
}

# Files/dirs to skip
SKIP_DIRS = {
    "__pycache__", ".git", ".ipynb_checkpoints", "node_modules",
    ".venv", "venv", ".tox", ".mypy_cache", ".pytest_cache",
}
SKIP_FILES = {".DS_Store", "Thumbs.db"}

# Max file size to include (500 KB)
MAX_FILE_SIZE = 500 * 1024


def should_include(filepath: str) -> bool:
    """Check if a file should be included based on extension and name."""
    basename = os.path.basename(filepath)
    if basename in SKIP_FILES:
        return False
    _, ext = os.path.splitext(basename)
    # Include Dockerfile (no extension)
    if basename.lower() in ("dockerfile", "makefile", "requirements.txt"):
        return True
    return ext.lower() in CODE_EXTENSIONS


def collect_files(directory: str, extra_skip_dirs: set[str] | None = None) -> list[tuple[str, str]]:
    """Walk directory and collect (relative_path, absolute_path) pairs."""
    skip = SKIP_DIRS | (extra_skip_dirs or set())
    files = []
    for root, dirs, filenames in os.walk(directory):
        # Skip hidden/unwanted directories
        dirs[:] = [d for d in dirs if d not in skip and not d.startswith(".")]
        dirs.sort()

        for fname in sorted(filenames):
            fpath = os.path.join(root, fname)
            if not should_include(fpath):
                continue
            try:
                if os.path.getsize(fpath) > MAX_FILE_SIZE:
                    continue
            except OSError:
                continue  # broken symlink or permission error
            rel = os.path.relpath(fpath, directory)
            files.append((rel, fpath))
    return files


def _add_code_block(doc, code: str):
    """Add a code block as a light-grey-background paragraph with monospace font."""
    for line in code.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(13)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        # Set monospace font for non-ASCII fallback
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.makeelement(qn("w:rFonts"), {
            qn("w:ascii"): "Consolas",
            qn("w:hAnsi"): "Consolas",
            qn("w:cs"): "Courier New",
        })
        rPr.insert(0, rFonts)


def _read_notebook_cells(filepath: str) -> list[dict]:
    """Parse a Jupyter notebook and return a list of cell dicts.

    Each dict has keys: cell_number (int), cell_type (str), source (str),
    and optionally outputs (list of text strings).
    """
    try:
        with open(filepath, encoding="utf-8", errors="replace") as f:
            nb = json.load(f)
    except (json.JSONDecodeError, ValueError) as e:
        return [{"cell_number": 1, "cell_type": "error", "source": f"[Could not parse notebook: {e}]"}]

    raw_cells = nb.get("cells", [])
    if not isinstance(raw_cells, list):
        return [{"cell_number": 1, "cell_type": "error", "source": "[Invalid notebook format]"}]

    cells = []
    for i, cell in enumerate(raw_cells, 1):
        cell_type = cell.get("cell_type", "code")
        source_raw = cell.get("source", [])
        source = "".join(source_raw) if isinstance(source_raw, list) else str(source_raw)

        # Extract text outputs (stdout, stderr, plain text results)
        output_texts = []
        for out in cell.get("outputs", []):
            if out.get("output_type") == "stream":
                text = "".join(out.get("text", []))
                output_texts.append(text)
            elif out.get("output_type") in ("execute_result", "display_data"):
                text_data = out.get("data", {}).get("text/plain", [])
                if isinstance(text_data, list):
                    output_texts.append("".join(text_data))
                elif isinstance(text_data, str):
                    output_texts.append(text_data)
            elif out.get("output_type") == "error":
                tb = out.get("traceback", [])
                # Strip ANSI escape codes from traceback
                clean = re.sub(r"\x1b\[[0-9;]*m", "", "\n".join(tb))
                output_texts.append(clean)

        cells.append({
            "cell_number": i,
            "cell_type": cell_type,
            "source": source,
            "outputs": output_texts,
        })
    return cells


def _add_notebook(doc, filepath: str):
    """Add a Jupyter notebook to the document with proper formatting for each cell type."""
    cells = _read_notebook_cells(filepath)

    for cell in cells:
        num = cell["cell_number"]
        ctype = cell["cell_type"]
        source = cell["source"]

        # Cell header
        header_p = doc.add_paragraph()
        header_p.paragraph_format.space_before = Pt(8)
        header_p.paragraph_format.space_after = Pt(2)
        run = header_p.add_run(f"Cell {num}  [{ctype}]")
        run.font.size = Pt(9)
        run.font.bold = True
        if ctype == "markdown":
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)  # blue for markdown
        elif ctype == "code":
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)  # dark for code
        else:
            run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)  # orange for raw/error

        # Cell content
        if ctype == "markdown":
            # Render markdown as regular text (preserving line breaks)
            for line in source.splitlines():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = Pt(14)
                # Handle markdown headers
                if line.startswith("#"):
                    stripped = line.lstrip("#").strip()
                    level = min(len(line) - len(line.lstrip("#")), 4)
                    run = p.add_run(stripped)
                    run.font.bold = True
                    run.font.size = Pt(12 - level)  # h1=11, h2=10, h3=9, h4=8
                else:
                    run = p.add_run(line if line else " ")
                    run.font.size = Pt(9)
        else:
            # Code and raw cells — monospace
            _add_code_block(doc, source)

        # Outputs
        outputs = cell.get("outputs", [])
        if outputs:
            out_header = doc.add_paragraph()
            out_header.paragraph_format.space_before = Pt(4)
            out_header.paragraph_format.space_after = Pt(1)
            run = out_header.add_run("Output:")
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)  # green

            combined = "\n".join(outputs)
            # Truncate very long outputs
            max_output_lines = 50
            lines = combined.splitlines()
            if len(lines) > max_output_lines:
                combined = "\n".join(lines[:max_output_lines]) + f"\n... ({len(lines) - max_output_lines} more lines truncated)"
            _add_code_block(doc, combined)


def generate_docx(directory: str, team_name: str, output_path: str, extra_skip_dirs: set[str] | None = None):
    """Generate the Word document."""
    files = collect_files(directory, extra_skip_dirs)
    if not files:
        print(f"No code files found in {directory}")
        sys.exit(1)

    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # -- Title page --
    doc.add_paragraph()  # spacer
    title = doc.add_heading(f"Code Submission", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(team_name)
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xED, 0x1C, 0x24)  # AMD red

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("TCS & AMD AI Hackathon 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # File count summary
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary.add_run(f"\n{len(files)} file(s) included")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # -- Table of contents --
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    for i, (rel_path, _) in enumerate(files, 1):
        p = doc.add_paragraph(f"{i}. {rel_path}", style="List Number")
        p.paragraph_format.space_after = Pt(2)

    # -- Each file --
    for i, (rel_path, abs_path) in enumerate(files, 1):
        doc.add_page_break()

        # File header
        heading = doc.add_heading(f"{rel_path}", level=2)

        # File metadata
        size = os.path.getsize(abs_path)
        size_str = f"{size:,} bytes" if size < 1024 else f"{size / 1024:.1f} KB"
        meta = doc.add_paragraph()
        run = meta.add_run(f"File {i} of {len(files)}  |  {size_str}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.italic = True

        # Separator
        doc.add_paragraph("_" * 80).runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # File content
        try:
            if rel_path.endswith(".ipynb"):
                _add_notebook(doc, abs_path)
            else:
                with open(abs_path, encoding="utf-8", errors="replace") as f:
                    content = f.read()
                _add_code_block(doc, content)
        except Exception as e:
            _add_code_block(doc, f"[Error reading file: {e}]")

    doc.save(output_path)
    print(f"Generated: {output_path} ({len(files)} files, {os.path.getsize(output_path) / 1024:.0f} KB)")


def main():
    parser = argparse.ArgumentParser(description="Generate a Word document from code files.")
    parser.add_argument("directory", help="Directory containing code files")
    parser.add_argument("--team", default="Team", help="Team name for the title page")
    parser.add_argument("--output", "-o", default=None, help="Output .docx path (default: <team>.docx)")
    parser.add_argument("--ignore-dirs", default=None, help="Comma-separated list of directory names to skip (e.g. data,logs,weights)")
    args = parser.parse_args()

    if not os.path.isdir(args.directory):
        print(f"Error: {args.directory} is not a directory")
        sys.exit(1)

    extra_skip = None
    if args.ignore_dirs:
        extra_skip = {d.strip() for d in args.ignore_dirs.split(",") if d.strip()}

    output = args.output or f"{args.team.replace(' ', '_')}.docx"
    if not output.endswith(".docx"):
        output += ".docx"

    generate_docx(args.directory, args.team, output, extra_skip)


if __name__ == "__main__":
    main()
